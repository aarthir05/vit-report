"""
Generate VIT Internship Report DOCX from content file.
Uses python-docx with VIT thesis formatting standards.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "AISPM_Internship_Report.docx")

def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def setup_styles(doc):
    """Configure document styles for VIT thesis format."""
    # Default style - Times New Roman, 12pt, 1.5 spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    # Heading 1 - Chapter headings (14pt, Bold, Centered)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.5

    # Heading 2 - Section headings (12pt, Bold)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.5

    # Heading 3 - Subsection headings (12pt, Bold Italic)
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.5

    return doc


def set_margins(doc):
    """Set VIT thesis margins: Left 1.5in (binding), others 1in."""
    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)


def add_page_numbers(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        
        run3 = paragraph.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)


def add_cover_page(doc):
    """Add the cover page."""
    # Add some spacing at top
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A project report on")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DESIGN AND DEVELOPMENT OF AN ENTERPRISE-GRADE AI SECURITY POSTURE MANAGEMENT PLATFORM FOR COMPREHENSIVE AI/ML ASSET GOVERNANCE AND RISK MITIGATION")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted in partial fulfillment for the award of the degree of")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bachelor of Technology")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("by")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[YOUR NAME] ([YOUR REG. NO.])")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("School of Computer Science and Engineering")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Vellore Institute of Technology, Vellore")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("June, 2026")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()


def add_declaration(doc):
    """Add Declaration page."""
    doc.add_heading('DECLARATION', level=1)
    
    doc.add_paragraph(
        'I hereby declare that the thesis entitled "DESIGN AND DEVELOPMENT OF AN '
        'ENTERPRISE-GRADE AI SECURITY POSTURE MANAGEMENT PLATFORM FOR COMPREHENSIVE '
        'AI/ML ASSET GOVERNANCE AND RISK MITIGATION" submitted by me, for the award of '
        'the degree of Bachelor of Technology in Computer Science and Engineering to '
        'Vellore Institute of Technology is a record of bonafide work carried out by me '
        'under the supervision of [Guide Name], [Designation], School of Computer Science '
        'and Engineering, VIT Vellore.'
    )
    
    doc.add_paragraph(
        'I further declare that the work reported in this thesis has not been submitted '
        'and will not be submitted, either in part or in full, for the award of any other '
        'degree or diploma in this institute or any other institute or university.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph('Place: Vellore')
    p = doc.add_paragraph('Date:')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Signature of the Candidate')
    
    doc.add_page_break()


def add_certificates(doc):
    """Add Certificate pages."""
    doc.add_heading('CERTIFICATE', level=1)
    doc.add_paragraph('[To be filled by your VIT internal guide]')
    doc.add_page_break()
    
    doc.add_heading('CERTIFICATE', level=1)
    doc.add_paragraph(
        'This is to certify that the project work entitled "Design and Development of an '
        'Enterprise-Grade AI Security Posture Management Platform for Comprehensive AI/ML '
        'Asset Governance and Risk Mitigation" is a bonafide record of work done by '
        '[YOUR NAME] (Reg. No. [YOUR REG NO]) during the period of internship at '
        'Ernst & Young (EY), under my supervision and guidance.'
    )
    doc.add_paragraph()
    doc.add_paragraph('[External Guide Name]')
    doc.add_paragraph('[Designation]')
    doc.add_paragraph('Ernst & Young LLP')
    doc.add_paragraph('[Location]')
    doc.add_paragraph()
    doc.add_paragraph('Date:')
    doc.add_page_break()


def add_abstract(doc):
    """Add Abstract page."""
    doc.add_heading('ABSTRACT', level=1)
    
    doc.add_paragraph(
        'The rapid proliferation of Artificial Intelligence and Machine Learning systems '
        'across enterprise environments has introduced a new category of security '
        'challenges that traditional cybersecurity tools are ill-equipped to address. '
        'Organizations deploying AI systems face unique threats including adversarial '
        'attacks, model poisoning, data leakage through inference, supply chain '
        'compromises in model artifacts, and regulatory non-compliance with emerging '
        'AI-specific legislation such as the EU AI Act, NIST AI Risk Management Framework, '
        'and ISO 42001.'
    )
    
    doc.add_paragraph(
        'This project presents the design and implementation of an AI Security Posture '
        'Management (AISPM) Platform \u2014 an enterprise-grade solution developed during an '
        'internship at Ernst & Young (EY) that provides comprehensive visibility, risk '
        'assessment, compliance monitoring, and automated governance for an organization\u2019s '
        'entire AI ecosystem. The platform encompasses 12 core modules addressing asset '
        'discovery, vulnerability management, LLM governance, responsible AI, data risk, '
        'vendor risk management, regulatory intelligence, and executive reporting.'
    )
    
    doc.add_paragraph(
        'The system architecture employs a modular monolith pattern built on FastAPI '
        '(Python 3.11) for the backend with 70+ specialized microservices, React 18 with '
        'TypeScript for the frontend comprising 28+ interactive pages, and a polyglot '
        'persistence layer utilizing PostgreSQL, Redis, Neo4j, and MongoDB. The platform '
        'integrates 14 industry-leading open-source tools including IBM\u2019s Adversarial '
        'Robustness Toolbox for attack simulation, Microsoft\u2019s Presidio for PII detection, '
        'Flower for federated learning, and OpenDP for differential privacy guarantees.'
    )
    
    doc.add_paragraph(
        'Key contributions include the development of an industry-first AI Bill of '
        'Materials (AI-BOM) generator compliant with CycloneDX 1.5, automated LLM safety '
        'testing covering prompt injection, hallucination, and data leakage vulnerabilities, '
        'and a multi-scanner asset discovery engine supporting 10 distinct source types '
        'across cloud, on-premise, and hybrid environments.'
    )
    
    doc.add_paragraph(
        'The platform achieves production-ready status with 26,000+ lines of code, '
        'comprehensive JWT-based authentication with role-based access control supporting '
        '7 user roles and 20+ granular permissions, Redis-backed distributed rate limiting, '
        'and full compliance coverage across EU AI Act, NIST AI RMF, ISO 42001, OWASP '
        'Top 10 for LLM, and MITRE ATLAS frameworks. Performance benchmarks demonstrate '
        'sub-2ms authentication overhead and p95 API response times under 200ms.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('Keywords: ')
    run.bold = True
    p.add_run(
        'AI Security, AISPM, Machine Learning Security, LLM Safety, AI '
        'Governance, Risk Management, Compliance Automation, FastAPI, React, Enterprise '
        'Security Platform'
    )
    
    doc.add_page_break()


def add_acknowledgement(doc):
    """Add Acknowledgement page."""
    doc.add_heading('ACKNOWLEDGEMENT', level=1)
    
    doc.add_paragraph(
        'It is my pleasure to express with deep sense of gratitude to [Internal Guide Name], '
        '[Designation], School of Computer Science and Engineering, Vellore Institute of '
        'Technology, for his/her constant guidance, continual encouragement, and '
        'understanding throughout this project. My association with him/her has been '
        'invaluable in shaping both the technical direction and academic rigor of this work.'
    )
    
    doc.add_paragraph(
        'I would like to express my sincere gratitude to my external guide at Ernst & Young, '
        '[External Guide Name], [Designation], for providing the opportunity to work on '
        'this cutting-edge project in AI security and for the mentorship that enabled me '
        'to develop industry-relevant skills in enterprise software development.'
    )
    
    doc.add_paragraph(
        'I would like to express my gratitude to the Chancellor, Vice Presidents, '
        'Vice Chancellor, Pro-Vice Chancellor, and [Dean Name], School of Computer Science '
        'and Engineering, for providing an exceptional academic environment and the '
        'infrastructure necessary to pursue this research.'
    )
    
    doc.add_paragraph(
        'I extend my heartfelt thanks to the [Program Chair Name], Program Chair and all '
        'teaching staff members of the School of Computer Science and Engineering for their '
        'guidance and encouragement throughout my academic journey at VIT.'
    )
    
    doc.add_paragraph(
        'I am deeply grateful to the EY AI Security team for their collaborative spirit, '
        'technical insights, and for creating an inclusive work environment that fostered '
        'innovation and learning. The exposure to enterprise-scale software development '
        'practices and cutting-edge AI security research has been transformative.'
    )
    
    doc.add_paragraph(
        'I would like to thank my parents for their unwavering support, patience, and '
        'encouragement throughout my education and internship period.'
    )
    
    doc.add_paragraph(
        'Finally, I express my appreciation to my friends and colleagues who provided '
        'valuable feedback, engaged in technical discussions, and motivated me to '
        'deliver my best work on this project.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('Place: Vellore')
    doc.add_paragraph('Date:')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('[YOUR NAME]')
    
    doc.add_page_break()


def add_table_of_contents(doc):
    """Add Table of Contents."""
    doc.add_heading('TABLE OF CONTENTS', level=1)
    
    toc_entries = [
        ("Abstract", "i"),
        ("Acknowledgement", "ii"),
        ("Table of Contents", "iii"),
        ("List of Tables", "vi"),
        ("List of Figures", "vii"),
        ("List of Abbreviations", "viii"),
        ("", ""),
        ("Chapter 1: Introduction", "1"),
        ("    1.1 Background and Motivation", "1"),
        ("    1.2 Problem Statement", "3"),
        ("    1.3 Objectives of the Project", "5"),
        ("    1.4 Scope of the Project", "6"),
        ("    1.5 Organization of the Report", "7"),
        ("", ""),
        ("Chapter 2: Literature Review", "8"),
        ("    2.1 AI Security Landscape", "8"),
        ("    2.2 Existing AI Security Frameworks", "10"),
        ("    2.3 AI-Specific Threat Taxonomy", "13"),
        ("    2.4 Regulatory Landscape for AI Systems", "16"),
        ("    2.5 Existing AISPM Solutions", "19"),
        ("    2.6 Open-Source Tools for AI Security", "21"),
        ("    2.7 Research Gaps and Motivation", "23"),
        ("", ""),
        ("Chapter 3: System Design and Architecture", "25"),
        ("    3.1 System Requirements", "25"),
        ("    3.2 High-Level Architecture", "27"),
        ("    3.3 Backend Architecture", "29"),
        ("    3.4 Frontend Architecture", "32"),
        ("    3.5 Database Design", "34"),
        ("    3.6 Security Architecture", "37"),
        ("    3.7 Deployment Architecture", "39"),
        ("", ""),
        ("Chapter 4: Implementation", "41"),
        ("    4.1 Development Environment Setup", "41"),
        ("    4.2 Backend Implementation", "42"),
        ("    4.3 Frontend Implementation", "48"),
        ("    4.4 Core Module Implementation", "51"),
        ("    4.5 LLM Safety and Red Teaming Implementation", "55"),
        ("    4.6 AI-BOM Generation Implementation", "57"),
        ("    4.7 Authentication and Authorization", "59"),
        ("", ""),
        ("Chapter 5: Testing and Results", "61"),
        ("    5.1 Testing Strategy", "61"),
        ("    5.2 Unit Testing", "62"),
        ("    5.3 Integration Testing", "63"),
        ("    5.4 Security Testing", "64"),
        ("    5.5 Performance Testing", "65"),
        ("    5.6 Results and Analysis", "66"),
        ("", ""),
        ("Chapter 6: Conclusion and Future Work", "68"),
        ("    6.1 Summary of Contributions", "68"),
        ("    6.2 Key Achievements", "69"),
        ("    6.3 Limitations", "70"),
        ("    6.4 Future Work", "70"),
        ("    6.5 Lessons Learned", "71"),
        ("", ""),
        ("References", "72"),
        ("Appendix A: API Endpoint Reference", "76"),
        ("Appendix B: Database Schema", "78"),
        ("Appendix C: Configuration Parameters", "80"),
    ]
    
    for entry, page in toc_entries:
        if not entry:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        if entry.startswith("    "):
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(entry.strip())
        elif entry.startswith("Chapter") or entry in ("References", "Appendix A: API Endpoint Reference", "Appendix B: Database Schema", "Appendix C: Configuration Parameters"):
            run = p.add_run(entry)
            run.bold = True
        else:
            p.add_run(entry)
        # Add tab and page number
        if page:
            p.add_run('\t' + page)
    
    doc.add_page_break()


def add_list_of_tables(doc):
    """Add List of Tables."""
    doc.add_heading('LIST OF TABLES', level=1)
    
    tables = [
        ("Table 2.1", "Comparison of AI Security Frameworks", "12"),
        ("Table 2.2", "OWASP Top 10 for LLM Applications", "14"),
        ("Table 2.3", "MITRE ATLAS Attack Techniques", "15"),
        ("Table 2.4", "Global AI Regulatory Landscape", "18"),
        ("Table 2.5", "Comparison of Existing AISPM Solutions", "20"),
        ("Table 3.1", "Functional Requirements", "25"),
        ("Table 3.2", "Non-Functional Requirements", "26"),
        ("Table 3.3", "Technology Stack Summary", "28"),
        ("Table 3.4", "Database Table Summary", "35"),
        ("Table 3.5", "User Roles and Permissions Matrix", "38"),
        ("Table 4.1", "Backend Service Modules", "43"),
        ("Table 4.2", "Frontend Page Components", "49"),
        ("Table 4.3", "Open-Source Integration Summary", "52"),
        ("Table 4.4", "LLM Safety Test Categories", "56"),
        ("Table 5.1", "Test Coverage Summary", "62"),
        ("Table 5.2", "Performance Benchmark Results", "66"),
        ("Table 5.3", "Security Test Results", "67"),
        ("Table 6.1", "Project Metrics Summary", "69"),
    ]
    
    for num, title, page in tables:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title}")
        p.add_run(f"\t{page}")
    
    doc.add_page_break()


def add_list_of_figures(doc):
    """Add List of Figures."""
    doc.add_heading('LIST OF FIGURES', level=1)
    
    figures = [
        ("Figure 1.1", "Growth of AI/ML Deployments in Enterprises (2020-2026)", "2"),
        ("Figure 1.2", "AI Security Incident Trends", "3"),
        ("Figure 2.1", "AI Threat Landscape Overview", "9"),
        ("Figure 2.2", "NIST AI Risk Management Framework", "11"),
        ("Figure 2.3", "EU AI Act Risk Classification", "17"),
        ("Figure 3.1", "High-Level System Architecture", "27"),
        ("Figure 3.2", "Backend Service Layer Architecture", "30"),
        ("Figure 3.3", "Frontend Component Hierarchy", "33"),
        ("Figure 3.4", "Entity-Relationship Diagram", "34"),
        ("Figure 3.5", "Authentication Flow Diagram", "37"),
        ("Figure 3.6", "Docker Compose Deployment Diagram", "39"),
        ("Figure 3.7", "Kubernetes Production Deployment", "40"),
        ("Figure 4.1", "Asset Discovery Flow", "44"),
        ("Figure 4.2", "Risk Assessment Pipeline", "46"),
        ("Figure 4.3", "AI-BOM Generation Workflow", "57"),
        ("Figure 4.4", "LLM Safety Testing Architecture", "55"),
        ("Figure 4.5", "JWT Authentication Sequence Diagram", "59"),
        ("Figure 5.1", "Test Pyramid Strategy", "61"),
        ("Figure 5.2", "API Response Time Distribution", "65"),
        ("Figure 5.3", "System Load Test Results", "66"),
    ]
    
    for num, title, page in figures:
        p = doc.add_paragraph()
        p.add_run(f"{num}  {title}")
        p.add_run(f"\t{page}")
    
    doc.add_page_break()


def add_abbreviations(doc):
    """Add List of Abbreviations."""
    doc.add_heading('LIST OF ABBREVIATIONS', level=1)
    
    abbreviations = [
        ("AISPM", "AI Security Posture Management"),
        ("AI", "Artificial Intelligence"),
        ("ML", "Machine Learning"),
        ("LLM", "Large Language Model"),
        ("API", "Application Programming Interface"),
        ("REST", "Representational State Transfer"),
        ("JWT", "JSON Web Token"),
        ("RBAC", "Role-Based Access Control"),
        ("CORS", "Cross-Origin Resource Sharing"),
        ("CSRF", "Cross-Site Request Forgery"),
        ("XSS", "Cross-Site Scripting"),
        ("SQL", "Structured Query Language"),
        ("ORM", "Object-Relational Mapping"),
        ("SPA", "Single Page Application"),
        ("CI/CD", "Continuous Integration/Continuous Deployment"),
        ("K8s", "Kubernetes"),
        ("FAIR", "Factor Analysis of Information Risk"),
        ("BOM", "Bill of Materials"),
        ("AI-BOM", "AI Bill of Materials"),
        ("CycloneDX", "OWASP CycloneDX Standard"),
        ("MITRE ATLAS", "Adversarial Threat Landscape for AI Systems"),
        ("OWASP", "Open Web Application Security Project"),
        ("NIST", "National Institute of Standards and Technology"),
        ("ISO", "International Organization for Standardization"),
        ("EU", "European Union"),
        ("GDPR", "General Data Protection Regulation"),
        ("HIPAA", "Health Insurance Portability and Accountability Act"),
        ("PII", "Personally Identifiable Information"),
        ("MFA", "Multi-Factor Authentication"),
        ("TOTP", "Time-based One-Time Password"),
        ("SSO", "Single Sign-On"),
        ("SAML", "Security Assertion Markup Language"),
        ("OAuth", "Open Authorization"),
        ("OIDC", "OpenID Connect"),
        ("CVE", "Common Vulnerabilities and Exposures"),
        ("SIEM", "Security Information and Event Management"),
        ("NLP", "Natural Language Processing"),
        ("SHAP", "SHapley Additive exPlanations"),
        ("LIME", "Local Interpretable Model-agnostic Explanations"),
        ("ART", "Adversarial Robustness Toolbox"),
        ("NYC LL144", "New York City Local Law 144"),
        ("BIPA", "Biometric Information Privacy Act"),
        ("DoS", "Denial of Service"),
        ("TLS", "Transport Layer Security"),
        ("AES", "Advanced Encryption Standard"),
        ("RSA", "Rivest-Shamir-Adleman"),
        ("HMAC", "Hash-based Message Authentication Code"),
        ("SLA", "Service Level Agreement"),
        ("KPI", "Key Performance Indicator"),
        ("MSP", "Managed Service Provider"),
        ("LOC", "Lines of Code"),
    ]
    
    # Create a table for abbreviations
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Abbreviation'
    hdr_cells[1].text = 'Full Form'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for abbr, full in abbreviations:
        row_cells = table.add_row().cells
        row_cells[0].text = abbr
        row_cells[1].text = full
    
    doc.add_page_break()


def add_chapter1(doc):
    """Add Chapter 1: Introduction."""
    doc.add_heading('CHAPTER 1: INTRODUCTION', level=1)
    
    # 1.1
    doc.add_heading('1.1 Background and Motivation', level=2)
    
    doc.add_paragraph(
        'The landscape of enterprise technology has undergone a paradigm shift with the '
        'widespread adoption of Artificial Intelligence (AI) and Machine Learning (ML) '
        'systems. According to McKinsey\u2019s 2025 Global AI Survey, over 72% of organizations '
        'have deployed at least one AI application in production, with the average '
        'enterprise maintaining 35-50 distinct AI models across their operations. This '
        'exponential growth in AI adoption has created an unprecedented challenge: how do '
        'organizations secure, govern, and maintain compliance for their rapidly expanding '
        'AI ecosystems?'
    )
    
    doc.add_paragraph(
        'Traditional cybersecurity tools and frameworks were designed for conventional '
        'software systems and are fundamentally inadequate for addressing AI-specific '
        'threats. Unlike traditional applications, AI systems are vulnerable to unique '
        'attack vectors including adversarial examples that manipulate model predictions, '
        'data poisoning attacks that corrupt training data, model extraction attacks that '
        'steal intellectual property, and inference attacks that compromise user privacy. '
        'The 2024 MITRE ATLAS framework documents over 60 distinct techniques that '
        'adversaries can employ against ML systems, representing a threat surface that '
        'conventional security tools cannot detect or mitigate.'
    )
    
    doc.add_paragraph(
        'The emergence of Large Language Models (LLMs) has further complicated the security '
        'landscape. Organizations deploying LLMs face novel threats such as prompt injection '
        'attacks, jailbreak vulnerabilities, hallucination risks, and data leakage through '
        'model outputs. The OWASP Top 10 for LLM Applications, published in 2023, '
        'highlights these critical vulnerabilities, yet most organizations lack the tooling '
        'to systematically test for and mitigate these risks.'
    )
    
    doc.add_paragraph(
        'Simultaneously, the regulatory environment for AI has intensified dramatically. '
        'The European Union\u2019s AI Act, which entered into force in 2024, mandates '
        'comprehensive governance requirements for AI systems based on risk classification. '
        'The NIST AI Risk Management Framework provides guidance for organizations in the '
        'United States, while ISO 42001 establishes an international standard for AI '
        'management systems. At the state level, regulations such as New York City\u2019s '
        'Local Law 144, Colorado\u2019s SB205, and Illinois\u2019 Biometric Information Privacy Act '
        'impose specific compliance obligations. Organizations operating globally must '
        'navigate this complex web of regulations while maintaining operational efficiency.'
    )
    
    doc.add_paragraph(
        'This confluence of factors \u2014 expanding AI deployments, novel threat vectors, and '
        'regulatory obligations \u2014 has given rise to the concept of AI Security Posture '
        'Management (AISPM). AISPM platforms provide unified visibility into an '
        'organization\u2019s AI assets, continuously assess security risks, monitor compliance '
        'posture, and enable automated governance workflows. However, the market for AISPM '
        'solutions is nascent, with existing tools offering fragmented capabilities that '
        'fail to address the full spectrum of AI security challenges.'
    )
    
    doc.add_paragraph(
        'This project was undertaken during an internship at Ernst & Young (EY), one of '
        'the world\u2019s largest professional services organizations, with the objective of '
        'designing and developing a comprehensive AISPM platform that addresses the '
        'complete lifecycle of AI security management. The platform aims to provide '
        'organizations with a single pane of glass for discovering all AI assets, assessing '
        'their security posture, ensuring regulatory compliance, and implementing automated '
        'governance controls.'
    )
    
    doc.add_paragraph('The motivation for this project stems from three critical observations:')
    
    doc.add_paragraph(
        'First, the gap between AI adoption speed and AI security maturity continues to '
        'widen. Organizations are deploying AI systems faster than they can secure them, '
        'creating a growing surface of unmanaged risk. A 2025 Gartner report indicates '
        'that fewer than 15% of organizations have implemented comprehensive AI security '
        'programs, despite 72% having production AI deployments.'
    )
    
    doc.add_paragraph(
        'Second, existing security tools provide inadequate coverage for AI-specific '
        'threats. Traditional vulnerability scanners, SIEM systems, and endpoint protection '
        'tools cannot detect adversarial attacks on ML models, identify biased algorithms, '
        'or assess compliance with AI-specific regulations. This gap necessitates '
        'purpose-built tooling designed specifically for the AI security domain.'
    )
    
    doc.add_paragraph(
        'Third, the fragmentation of the current AISPM market means organizations must '
        'assemble multiple point solutions to achieve comprehensive coverage, leading to '
        'integration complexity, visibility gaps, and increased total cost of ownership. '
        'A unified platform that consolidates AI security, governance, and compliance '
        'capabilities represents a significant advancement for enterprise security teams.'
    )
    
    # 1.2
    doc.add_heading('1.2 Problem Statement', level=2)
    
    doc.add_paragraph(
        'Organizations deploying AI/ML systems at scale face a multifaceted security and '
        'governance challenge characterized by the following problems:'
    )
    
    problems = [
        ('1. Lack of Visibility:', 'Most organizations cannot enumerate all AI assets deployed across their environment. Shadow AI \u2014 unauthorized use of AI services by employees \u2014 creates blind spots in the security posture. Models may be deployed in cloud environments, Kubernetes clusters, code repositories, or accessed via third-party APIs without centralized tracking or governance oversight.'),
        ('2. AI-Specific Vulnerabilities:', 'Traditional vulnerability management tools cannot identify AI-specific attack surfaces. Threats such as adversarial examples, model backdoors, data poisoning, membership inference attacks, and model extraction require specialized detection and assessment capabilities that do not exist in conventional security stacks.'),
        ('3. LLM Security Gaps:', 'The rapid deployment of Large Language Models introduces novel vulnerability classes including prompt injection, jailbreak attacks, hallucination risks, and potential data leakage through model outputs. Organizations lack systematic testing frameworks to evaluate LLM security before and during deployment.'),
        ('4. Regulatory Complexity:', 'The proliferating landscape of AI regulations across jurisdictions creates compliance burdens that organizations struggle to manage manually. Different frameworks (EU AI Act, NIST AI RMF, ISO 42001, state-level regulations) have overlapping yet distinct requirements, making comprehensive compliance assessment a complex undertaking.'),
        ('5. Supply Chain Risk:', 'AI model supply chains introduce unique risks through pre-trained models, third-party datasets, and open-source ML frameworks that may contain vulnerabilities, backdoors, or biases. Unlike software supply chain security, AI supply chain risk lacks standardized assessment methodologies and tooling.'),
        ('6. Governance Fragmentation:', 'AI governance responsibilities are typically distributed across security teams, data science teams, compliance teams, and business units without a unifying framework or platform. This fragmentation leads to inconsistent policy enforcement, gaps in oversight, and inability to demonstrate governance to regulators and auditors.'),
    ]
    
    for title, desc in problems:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph(
        'The central research question addressed by this project is: How can a unified, '
        'enterprise-grade platform be designed and implemented to provide comprehensive '
        'AI security posture management, encompassing asset discovery, risk assessment, '
        'compliance monitoring, vulnerability management, and automated governance across '
        'an organization\u2019s entire AI ecosystem?'
    )
    
    # 1.3
    doc.add_heading('1.3 Objectives of the Project', level=2)
    
    doc.add_paragraph('The primary objectives of this project are:')
    
    objectives = [
        'Design and implement a comprehensive AI Security Posture Management platform that provides unified visibility into all AI/ML assets across cloud, on-premise, and hybrid environments.',
        'Develop a multi-scanner asset discovery engine capable of identifying AI assets across 10 distinct source types including cloud ML platforms, Kubernetes clusters, code repositories, file systems, and network traffic.',
        'Implement automated risk assessment using the FAIR (Factor Analysis of Information Risk) methodology, providing quantitative risk scoring specific to AI/ML threats.',
        'Build a comprehensive compliance monitoring engine supporting multiple regulatory frameworks including EU AI Act, NIST AI RMF, ISO 42001, OWASP Top 10 for LLM, and MITRE ATLAS.',
        'Develop an LLM safety testing framework capable of systematically evaluating prompt injection, hallucination, toxicity, and data leakage vulnerabilities.',
        'Create an industry-first AI Bill of Materials (AI-BOM) generator compliant with the CycloneDX 1.5 specification for AI supply chain transparency.',
        'Implement role-based access control with comprehensive audit logging to support enterprise governance requirements and regulatory compliance.',
        'Design and deploy the platform using modern DevOps practices including containerization, orchestration, CI/CD, and infrastructure monitoring.',
        'Integrate 14 open-source AI security tools into a cohesive platform that provides capabilities not available in any single commercial solution.',
        'Achieve production-ready status with enterprise-grade security, performance, and reliability characteristics.',
    ]
    
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph(f'{i}. {obj}')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 1.4
    doc.add_heading('1.4 Scope of the Project', level=2)
    
    doc.add_paragraph('The scope of this project encompasses:')
    
    p = doc.add_paragraph()
    run = p.add_run('In Scope:')
    run.bold = True
    
    in_scope = [
        'Backend API development using Python/FastAPI with 70+ service modules',
        'Frontend SPA development using React 18/TypeScript with 28+ pages',
        'Database design and implementation (PostgreSQL, Redis, Neo4j, MongoDB)',
        'Authentication and authorization system (JWT + RBAC)',
        'Asset discovery engine with 10 scanner types',
        'Risk assessment using FAIR methodology',
        'Compliance monitoring for 5+ regulatory frameworks',
        'LLM safety testing (5 vulnerability categories)',
        'AI-BOM generation (CycloneDX 1.5)',
        'Vendor risk management',
        'Executive reporting and dashboards',
        'Docker and Kubernetes deployment configurations',
        'CI/CD pipeline configuration',
        'Monitoring and alerting setup',
        'Comprehensive documentation (150+ pages)',
        'Testing (unit, integration, security, performance)',
    ]
    
    for item in in_scope:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('Out of Scope:')
    run.bold = True
    
    out_scope = [
        'Production deployment to live enterprise environments',
        'Integration with actual customer AI infrastructure',
        'Real-time network traffic analysis deployment',
        'Compliance certification (SOC 2, ISO 27001) auditing',
        'Mobile application full release',
        'Commercial licensing and distribution',
        'Customer onboarding and training delivery',
    ]
    
    for item in out_scope:
        doc.add_paragraph(item, style='List Bullet')
    
    # 1.5
    doc.add_heading('1.5 Organization of the Report', level=2)
    
    doc.add_paragraph('This report is organized into six chapters:')
    
    chapters = [
        ('Chapter 1 (Introduction)', 'provides the background, motivation, problem statement, objectives, and scope of the project.'),
        ('Chapter 2 (Literature Review)', 'examines the current AI security landscape, existing frameworks and tools, regulatory requirements, and identifies research gaps that motivate this work.'),
        ('Chapter 3 (System Design and Architecture)', 'presents the system requirements, architectural decisions, and detailed design of the backend, frontend, database, security, and deployment layers.'),
        ('Chapter 4 (Implementation)', 'describes the implementation details of core modules, including the asset discovery engine, LLM safety testing framework, AI-BOM generator, and authentication system.'),
        ('Chapter 5 (Testing and Results)', 'covers the testing strategy, methodologies employed, and presents results from unit, integration, security, and performance testing.'),
        ('Chapter 6 (Conclusion and Future Work)', 'summarizes contributions, discusses limitations, outlines future enhancement directions, and presents lessons learned.'),
    ]
    
    for title, desc in chapters:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_page_break()


def add_chapter2(doc):
    """Add Chapter 2: Literature Review."""
    doc.add_heading('CHAPTER 2: LITERATURE REVIEW', level=1)
    
    # 2.1
    doc.add_heading('2.1 AI Security Landscape', level=2)
    
    paragraphs_2_1 = [
        'The security of AI systems has emerged as a critical concern in both academic research and industry practice. Unlike traditional software systems where security focuses primarily on preventing unauthorized access and ensuring data integrity, AI systems present a fundamentally different attack surface rooted in the statistical nature of machine learning algorithms.',
        'Goodfellow et al. (2015) first demonstrated the existence of adversarial examples \u2014 carefully crafted inputs that cause ML models to produce incorrect outputs with high confidence. This seminal work revealed that the decision boundaries learned by neural networks are inherently fragile and can be exploited by adversaries with minimal perturbation. Subsequent research has expanded the taxonomy of adversarial attacks to include evasion attacks (modifying inputs at inference time), poisoning attacks (corrupting training data), model extraction (stealing model functionality), and inference attacks (extracting private information from model outputs).',
        'The AI security landscape can be categorized into four primary domains:',
    ]
    for p_text in paragraphs_2_1:
        doc.add_paragraph(p_text)
    
    domains = [
        ('Model Security:', 'Protecting the integrity, confidentiality, and availability of ML models. This encompasses defense against adversarial attacks, model watermarking for intellectual property protection, secure model serving, and detection of model tampering or backdoors (Gu et al., 2019).'),
        ('Data Security:', 'Ensuring the privacy and integrity of training data, inference data, and model outputs. Key concerns include training data poisoning (Biggio & Roli, 2018), membership inference attacks (Shokri et al., 2017), and model inversion attacks that reconstruct training data from model parameters (Fredrikson et al., 2015).'),
        ('Supply Chain Security:', 'Managing risks arising from dependencies on third-party models, datasets, and ML frameworks. Research by Goldblum et al. (2022) demonstrates that pre-trained models downloaded from public repositories may contain backdoors that activate on specific trigger inputs.'),
        ('Operational Security:', 'Monitoring deployed models for performance degradation, concept drift, and emergent behaviors. Sculley et al. (2015) identified the concept of \u201ctechnical debt\u201d in ML systems, highlighting how the operational complexity of AI systems exceeds that of traditional software.'),
    ]
    for title, desc in domains:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph(
        'The emergence of Large Language Models has introduced additional security dimensions. Perez and Ribeiro (2022) demonstrated prompt injection attacks where adversarial instructions embedded in input text override the model\u2019s intended behavior. Carlini et al. (2023) showed that LLMs can memorize and reproduce training data, creating privacy risks. These findings underscore the need for specialized security testing frameworks designed for the unique characteristics of generative AI systems.'
    )
    
    # 2.2
    doc.add_heading('2.2 Existing AI Security Frameworks', level=2)
    
    doc.add_paragraph('Several frameworks have been developed to guide organizations in managing AI security risks:')
    
    p = doc.add_paragraph()
    run = p.add_run('NIST AI Risk Management Framework (AI RMF):')
    run.bold = True
    doc.add_paragraph(
        'Released in January 2023, the NIST AI RMF provides a voluntary framework for managing risks throughout the AI lifecycle. It is organized around four core functions: Govern, Map, Measure, and Manage. The framework emphasizes the importance of organizational governance structures, risk identification and mapping, measurement of AI risks through metrics and monitoring, and management through mitigation strategies. While comprehensive in its guidance, the framework does not prescribe specific technical controls, leaving implementation decisions to organizations.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('MITRE ATLAS (Adversarial Threat Landscape for Artificial Intelligence Systems):')
    run.bold = True
    doc.add_paragraph(
        'MITRE ATLAS, launched in 2021, provides a knowledge base of adversarial techniques targeting AI systems, analogous to the MITRE ATT&CK framework for traditional cybersecurity. ATLAS documents attack techniques across 12 tactics including Reconnaissance, Resource Development, Initial Access, ML Model Access, ML Attack Staging, and ML Attack Execution. The framework includes over 60 techniques with real-world case studies, providing a structured taxonomy for threat modeling of AI systems.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('OWASP Top 10 for LLM Applications:')
    run.bold = True
    doc.add_paragraph('Published in August 2023, this framework identifies the ten most critical security risks for applications built on Large Language Models. The risks include:')
    
    owasp_items = [
        'LLM01 - Prompt Injection: Manipulation of LLM behavior through crafted inputs',
        'LLM02 - Insecure Output Handling: Insufficient validation of model outputs',
        'LLM03 - Training Data Poisoning: Corruption of training data to influence behavior',
        'LLM04 - Model Denial of Service: Resource exhaustion attacks against LLM endpoints',
        'LLM05 - Supply Chain Vulnerabilities: Risks from third-party models and plugins',
        'LLM06 - Sensitive Information Disclosure: Leakage of private data in responses',
        'LLM07 - Insecure Plugin Design: Vulnerabilities in LLM tool-use implementations',
        'LLM08 - Excessive Agency: LLMs performing actions beyond intended scope',
        'LLM09 - Overreliance: Insufficient human oversight of AI-generated content',
        'LLM10 - Model Theft: Unauthorized extraction of model weights or capabilities',
    ]
    for item in owasp_items:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('ISO/IEC 42001:2023 (AI Management System):')
    run.bold = True
    doc.add_paragraph(
        'This international standard specifies requirements for establishing, implementing, maintaining, and continually improving an AI management system. It provides a framework for organizations to manage AI-related risks and opportunities, with requirements covering leadership commitment, planning, support, operation, performance evaluation, and improvement. ISO 42001 takes a process-oriented approach, focusing on organizational governance rather than technical controls.'
    )
    
    p = doc.add_paragraph()
    run = p.add_run('EU AI Act Risk Classification:')
    run.bold = True
    doc.add_paragraph('The EU AI Act classifies AI systems into four risk categories:')
    
    risk_cats = [
        'Unacceptable Risk: Banned applications (social scoring, real-time biometric identification in public spaces)',
        'High Risk: Stringent requirements (employment decisions, credit scoring, critical infrastructure)',
        'Limited Risk: Transparency obligations (chatbots, deepfakes)',
        'Minimal Risk: No specific obligations',
    ]
    for item in risk_cats:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('Each category imposes distinct requirements ranging from complete prohibition to voluntary codes of practice, creating a graduated regulatory approach that organizations must navigate.')
    
    # 2.3
    doc.add_heading('2.3 AI-Specific Threat Taxonomy', level=2)
    
    doc.add_paragraph('AI systems face a unique threat landscape that extends beyond traditional cybersecurity concerns. This section categorizes AI-specific threats according to the stage of the AI lifecycle they target:')
    
    doc.add_heading('Training Phase Attacks:', level=3)
    
    p = doc.add_paragraph()
    run = p.add_run('Data Poisoning: ')
    run.bold = True
    p.add_run('Adversaries inject malicious samples into training data to influence model behavior. Biggio and Roli (2018) demonstrate that poisoning as little as 3% of training data can significantly degrade model accuracy or introduce targeted misclassifications. Backdoor attacks represent a particularly insidious form of poisoning where the model behaves normally on clean inputs but produces adversary-chosen outputs when a specific trigger pattern is present (Gu et al., 2019).')
    
    p = doc.add_paragraph()
    run = p.add_run('Label Flipping: ')
    run.bold = True
    p.add_run('A subset of poisoning attacks where adversaries corrupt the labels in training data rather than the features. This approach is more difficult to detect through data quality checks and can be executed by adversaries with access to annotation pipelines or crowdsourcing platforms.')
    
    doc.add_heading('Inference Phase Attacks:', level=3)
    
    p = doc.add_paragraph()
    run = p.add_run('Adversarial Examples: ')
    run.bold = True
    p.add_run('Inputs crafted to cause misclassification while being imperceptible to human observers. Attacks range from white-box methods requiring full model access (Carlini & Wagner, 2017) to black-box attacks using only model outputs (Chen et al., 2017). Physical-world adversarial attacks have been demonstrated against autonomous driving systems, facial recognition, and medical imaging classifiers.')
    
    p = doc.add_paragraph()
    run = p.add_run('Model Extraction: ')
    run.bold = True
    p.add_run('Query-based attacks that replicate model functionality through systematic probing. Tramer et al. (2016) showed that models deployed as APIs can be effectively stolen through a polynomial number of queries, enabling adversaries to obtain a functionally equivalent copy of proprietary models.')
    
    p = doc.add_paragraph()
    run = p.add_run('Membership Inference: ')
    run.bold = True
    p.add_run('Attacks that determine whether a specific data point was part of the model\u2019s training set. Shokri et al. (2017) demonstrated that this information can be inferred from model outputs alone, creating privacy violations particularly concerning for models trained on sensitive data.')
    
    p = doc.add_paragraph()
    run = p.add_run('Model Inversion: ')
    run.bold = True
    p.add_run('Attacks that reconstruct sensitive features of training data from model parameters or predictions. Fredrikson et al. (2015) showed that facial recognition models can be inverted to reconstruct training images.')
    
    doc.add_heading('LLM-Specific Threats:', level=3)
    
    llm_threats = [
        ('Prompt Injection:', 'Direct injection manipulates the model\u2019s behavior by embedding adversarial instructions in user input. Indirect injection embeds malicious instructions in content the LLM processes (web pages, documents, emails), enabling attacks even when the adversary cannot directly interact with the model.'),
        ('Jailbreaking:', 'Techniques that bypass safety guardrails through role-playing scenarios, encoding tricks, or adversarial prompt construction. Successful jailbreaks can cause models to generate harmful, biased, or policy-violating content.'),
        ('Hallucination:', 'LLMs generate plausible but factually incorrect information. In enterprise contexts, hallucinations in decision-support systems can lead to flawed business decisions, legal liability, or safety risks.'),
        ('Data Leakage:', 'LLMs trained on confidential data may reproduce that data in responses. Carlini et al. (2023) demonstrated extraction of memorized training data from GPT-2, including personally identifiable information, code snippets, and URLs.'),
    ]
    for title, desc in llm_threats:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_heading('Supply Chain Threats:', level=3)
    
    sc_threats = [
        ('Model Supply Chain:', 'Pre-trained models from public repositories (Hugging Face, PyTorch Hub) may contain trojans or backdoors. Goldblum et al. (2022) showed that backdoored models uploaded to model registries could compromise downstream applications.'),
        ('Framework Vulnerabilities:', 'ML frameworks (TensorFlow, PyTorch, scikit-learn) may contain security vulnerabilities. The deserialization of model files (pickle, joblib) is particularly dangerous as it enables arbitrary code execution.'),
        ('Dataset Supply Chain:', 'Public datasets may contain biases, mislabeled data, or intentionally corrupted samples that affect model behavior when used for training or fine-tuning.'),
    ]
    for title, desc in sc_threats:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    # 2.4
    doc.add_heading('2.4 Regulatory Landscape for AI Systems', level=2)
    
    doc.add_paragraph('The regulatory environment for AI has evolved rapidly since 2023, with multiple jurisdictions implementing AI-specific legislation:')
    
    p = doc.add_paragraph()
    run = p.add_run('European Union - AI Act (2024):')
    run.bold = True
    doc.add_paragraph('The EU AI Act represents the world\u2019s first comprehensive AI regulation. Key requirements for high-risk AI systems include:')
    
    eu_reqs = [
        'Risk management system implementation and maintenance',
        'Data governance for training, validation, and testing datasets',
        'Technical documentation and record-keeping',
        'Transparency and provision of information to users',
        'Human oversight measures',
        'Accuracy, robustness, and cybersecurity requirements',
        'Conformity assessment before market placement',
        'Post-market monitoring and incident reporting',
    ]
    for item in eu_reqs:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('Non-compliance penalties range up to 35 million euros or 7% of global annual turnover, making compliance a critical business imperative.')
    
    p = doc.add_paragraph()
    run = p.add_run('United States - NIST AI RMF and State Regulations:')
    run.bold = True
    doc.add_paragraph('While the US lacks a comprehensive federal AI law, the NIST AI Risk Management Framework provides voluntary guidance widely adopted by industry. Additionally, state-level regulations impose specific requirements:')
    
    doc.add_paragraph('New York City Local Law 144 (2023): Requires bias audits for automated employment decision tools, with annual independent audits and public disclosure of results.')
    doc.add_paragraph('Colorado SB21-169 and SB205: Establishes governance requirements for AI systems used in insurance underwriting and consumer-facing decisions, including algorithmic accountability measures.')
    doc.add_paragraph('Illinois Biometric Information Privacy Act (BIPA): Imposes strict consent requirements for collection and use of biometric data, with a private right of action enabling individual lawsuits for violations.')
    
    p = doc.add_paragraph()
    run = p.add_run('International Standards:')
    run.bold = True
    doc.add_paragraph('ISO/IEC 42001:2023 provides an internationally recognized standard for AI management systems, while ISO/IEC 23894:2023 offers guidance on AI risk management. These standards provide a framework for organizations seeking to demonstrate AI governance maturity to customers, partners, and regulators.')
    
    p = doc.add_paragraph()
    run = p.add_run('India - Digital Personal Data Protection Act (DPDPA):')
    run.bold = True
    doc.add_paragraph('India\u2019s 2023 data protection legislation includes provisions relevant to AI systems processing personal data, with requirements for purpose limitation, data minimization, and individual rights that AI systems must respect.')
    
    doc.add_paragraph('The convergence of these regulatory frameworks creates a complex compliance landscape where organizations operating globally must satisfy multiple overlapping requirements. This complexity motivates the development of automated compliance monitoring tools that can map AI systems against multiple regulatory frameworks simultaneously.')
    
    # 2.5
    doc.add_heading('2.5 Existing AISPM Solutions', level=2)
    
    doc.add_paragraph('The market for AI Security Posture Management solutions is nascent but growing rapidly. Key commercial solutions include:')
    
    solutions = [
        ('HolisticAI:', 'A UK-based company offering AI governance and risk management tools. Their platform provides bias auditing, compliance assessment (EU AI Act, ISO 42001), and AI risk scoring. Strengths include strong compliance coverage and executive reporting. Limitations include lack of technical security testing (no adversarial attack simulation), no LLM-specific safety testing, and limited asset discovery capabilities.'),
        ('Protect AI:', 'Focused specifically on ML security, Protect AI offers model scanning, supply chain security, and vulnerability detection for ML pipelines. Their ModelScan tool detects malicious code in serialized models. Strengths include strong technical security capabilities. Limitations include no compliance automation, limited governance features, and no LLM safety testing.'),
        ('Robust Intelligence:', 'Provides AI validation and security testing including adversarial robustness evaluation, data quality assessment, and model monitoring. Strengths include automated testing capabilities. Limitations include narrow focus on model validation without broader AISPM capabilities.'),
        ('Calypso AI:', 'Offers AI security testing with focus on LLM applications, including prompt injection and jailbreak testing. Strengths include LLM-focused security. Limitations include lack of asset discovery, compliance automation, and vendor risk management.'),
    ]
    for title, desc in solutions:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    p = doc.add_paragraph()
    run = p.add_run('Gap Analysis:')
    run.bold = True
    doc.add_paragraph('No existing solution provides the full breadth of capabilities required for comprehensive AI security posture management. The market is characterized by:')
    
    gaps = [
        'Point solutions addressing individual aspects of AI security',
        'Lack of integration between security, governance, and compliance tools',
        'Absence of unified asset discovery across diverse environments',
        'No solution offering all of: adversarial testing, LLM safety, compliance automation, vendor risk, and executive reporting in a single platform',
        'Limited open-source integration preventing customization',
    ]
    for item in gaps:
        doc.add_paragraph(item, style='List Bullet')
    
    # 2.6
    doc.add_heading('2.6 Open-Source Tools for AI Security', level=2)
    
    doc.add_paragraph('The open-source community has developed several tools addressing specific aspects of AI security:')
    
    tools = [
        ('Adversarial Robustness Toolbox (ART) - IBM Research:', 'A Python library providing tools for adversarial attack and defense in ML. ART supports all major ML frameworks and implements attacks including FGSM, PGD, C&W, and DeepFool, along with certified defenses and detection mechanisms.'),
        ('AI Fairness 360 (AIF360) - IBM Research:', 'A comprehensive toolkit for detecting and mitigating bias in ML models. AIF360 implements over 70 fairness metrics across pre-processing, in-processing, and post-processing stages.'),
        ('Fickling - Trail of Bits:', 'A Python library for analyzing and modifying pickle files, the most common serialization format for ML models. Fickling can detect malicious code injected into pickle files, preventing supply chain attacks through compromised model artifacts.'),
        ('Flower (FL):', 'A federated learning framework that enables training ML models across distributed datasets without centralizing data. Flower supports a variety of federation strategies and provides privacy guarantees through secure aggregation.'),
        ('OpenDP:', 'A library for differential privacy that provides mathematical guarantees about the privacy of individual data points used in analyses. OpenDP implements various noise mechanisms and composition theorems.'),
        ('Evidently AI:', 'An open-source ML monitoring framework that detects data drift, model performance degradation, and data quality issues.'),
        ('Great Expectations:', 'A data validation framework that enables definition of expectations about data quality and automatic verification against incoming data.'),
        ('Microsoft Presidio:', 'An open-source tool for PII detection and anonymization supporting multiple entities across multiple languages.'),
    ]
    for title, desc in tools:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    # 2.7
    doc.add_heading('2.7 Research Gaps and Motivation', level=2)
    
    doc.add_paragraph('The literature review reveals several significant gaps that motivate this project:')
    
    research_gaps = [
        ('1. Integration Gap:', 'No existing platform integrates the full spectrum of AI security capabilities (discovery, assessment, compliance, LLM safety, supply chain) into a unified solution.'),
        ('2. Discovery Gap:', 'Existing tools assume knowledge of AI assets rather than providing automated discovery. Organizations cannot secure what they cannot see.'),
        ('3. LLM + Traditional AI Gap:', 'Solutions tend to focus either on traditional ML security or LLM security. No platform addresses both in a unified framework.'),
        ('4. Compliance Automation Gap:', 'While frameworks like the EU AI Act prescribe requirements, automated tools for continuously assessing compliance against multiple frameworks simultaneously are lacking.'),
        ('5. AI-BOM Gap:', 'Despite the importance of supply chain transparency, no widely adopted tool generates comprehensive AI Bills of Materials following established standards like CycloneDX.'),
        ('6. Open-Source Integration Gap:', 'Individual open-source tools provide excellent capabilities but require significant engineering effort to integrate, configure, and operationalize.'),
    ]
    for title, desc in research_gaps:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph('This project addresses these gaps by designing and implementing a unified AISPM platform that integrates asset discovery, risk assessment, compliance monitoring, LLM safety testing, AI-BOM generation, and governance automation into a single enterprise-grade solution.')
    
    doc.add_page_break()


def add_chapter3(doc):
    """Add Chapter 3: System Design and Architecture."""
    doc.add_heading('CHAPTER 3: SYSTEM DESIGN AND ARCHITECTURE', level=1)
    
    # 3.1
    doc.add_heading('3.1 System Requirements', level=2)
    doc.add_heading('3.1.1 Functional Requirements', level=3)
    doc.add_paragraph('The platform must satisfy the following functional requirements:')
    
    func_reqs = [
        'FR1 - Asset Discovery: The system shall automatically discover AI/ML assets across cloud platforms (AWS SageMaker, Azure ML, Google Vertex AI), Kubernetes clusters, code repositories, file systems, MLOps platforms, and network traffic.',
        'FR2 - Asset Inventory: The system shall maintain a centralized inventory of all discovered AI assets with comprehensive metadata.',
        'FR3 - Risk Assessment: The system shall perform automated risk assessment using FAIR methodology, generating quantitative risk scores.',
        'FR4 - Compliance Monitoring: The system shall continuously monitor AI assets against regulatory frameworks including EU AI Act, NIST AI RMF, ISO 42001.',
        'FR5 - Vulnerability Management: The system shall identify and track AI-specific vulnerabilities mapped to MITRE ATLAS and OWASP Top 10 for LLM.',
        'FR6 - LLM Safety Testing: The system shall provide automated testing for LLM vulnerabilities including prompt injection, hallucination, toxicity, and data leakage.',
        'FR7 - AI-BOM Generation: The system shall generate AI Bills of Materials compliant with CycloneDX 1.5 specification.',
        'FR8 - Authentication and Authorization: The system shall implement secure authentication with role-based access control.',
        'FR9 - Vendor Risk Management: The system shall enable assessment and monitoring of third-party AI vendor security posture.',
        'FR10 - Reporting: The system shall generate automated audit reports, executive dashboards, and compliance documentation.',
        'FR11 - Shadow AI Detection: The system shall detect unauthorized AI service usage through network traffic analysis.',
        'FR12 - Real-time Monitoring: The system shall provide real-time dashboards with WebSocket-based updates.',
    ]
    for req in func_reqs:
        doc.add_paragraph(req)
    
    doc.add_heading('3.1.2 Non-Functional Requirements', level=3)
    
    nfr = [
        'NFR1 - Performance: API response time shall be less than 200ms at the 95th percentile under normal load conditions.',
        'NFR2 - Security: The system shall implement defense-in-depth security including encryption at rest (AES-256), encryption in transit (TLS 1.3), input validation, rate limiting, and security headers.',
        'NFR3 - Scalability: The system shall support horizontal scaling through containerization and Kubernetes orchestration, handling up to 1000 concurrent users.',
        'NFR4 - Availability: The system shall target 99.9% uptime through health checks, auto-restart capabilities, and load-balanced deployment.',
        'NFR5 - Maintainability: The system shall follow modular architecture with separation of concerns.',
        'NFR6 - Usability: The frontend shall provide an intuitive interface requiring minimal training.',
        'NFR7 - Extensibility: The system shall support addition of new scanner types, compliance frameworks, and integration connectors without core architecture changes.',
        'NFR8 - Compliance: The system itself shall comply with security best practices including OWASP Top 10.',
    ]
    for req in nfr:
        doc.add_paragraph(req)
    
    # 3.2
    doc.add_heading('3.2 High-Level Architecture', level=2)
    
    doc.add_paragraph('The AISPM platform employs a modular monolith architecture \u2014 a single deployable unit with clear internal module boundaries. This architectural choice balances the benefits of microservices (modularity, separation of concerns) with the operational simplicity of a monolith (single deployment, simplified data consistency, reduced latency for inter-module communication).')
    
    doc.add_paragraph('The system comprises four primary layers:')
    
    layers = [
        ('Presentation Layer:', 'React 18 SPA communicating with the backend via RESTful APIs and WebSocket connections for real-time updates. The EY Motif Design System provides visual consistency and accessibility compliance.'),
        ('Application Layer:', 'FastAPI backend with 70+ service modules organized by domain (security, compliance, governance, data, integration). Each module encapsulates its business logic, data access, and API endpoints while sharing common infrastructure.'),
        ('Data Layer:', 'Polyglot persistence utilizing PostgreSQL for transactional data, Redis for caching and rate limiting, Neo4j for relationship graphs and attack paths, and MongoDB for audit logs and document storage.'),
        ('Infrastructure Layer:', 'Docker containerization with Kubernetes orchestration, Nginx reverse proxy, Prometheus/Grafana monitoring, and GitHub Actions CI/CD.'),
    ]
    for title, desc in layers:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    # 3.3
    doc.add_heading('3.3 Backend Architecture', level=2)
    
    doc.add_paragraph('The backend is built on FastAPI, chosen for its combination of high performance (comparable to Node.js and Go), native async/await support, automatic API documentation generation, and type-safe request/response validation through Pydantic.')
    
    doc.add_heading('3.3.1 Application Structure', level=3)
    
    structure = [
        ('API Layer (Routers):', 'FastAPI router modules that define endpoint paths, request validation schemas, response models, and dependency injection.'),
        ('Service Layer:', 'Business logic encapsulated in service classes. Services coordinate between data access, external integrations, and business rules.'),
        ('Data Access Layer:', 'SQLAlchemy models and async session management providing abstracted database operations.'),
        ('Integration Layer:', 'Connectors to external systems with interface abstractions enabling testing and configuration flexibility.'),
    ]
    for title, desc in structure:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_heading('3.3.2 Middleware Pipeline', level=3)
    doc.add_paragraph('The FastAPI middleware pipeline processes every request through:')
    
    middleware = [
        'CORS validation (origin whitelisting)',
        'CSRF token verification (state-changing requests)',
        'Request size validation (DoS prevention)',
        'Rate limiting (Redis-backed, per-client)',
        'Security header injection (response)',
        'Authentication validation (JWT verification)',
        'Authorization check (RBAC permission verification)',
        'Audit logging (all authenticated actions)',
    ]
    for i, item in enumerate(middleware, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    # 3.4
    doc.add_heading('3.4 Frontend Architecture', level=2)
    doc.add_paragraph('The frontend employs React 18 with TypeScript, prioritizing type safety, component reusability, and developer productivity.')
    
    doc.add_heading('3.4.1 Component Architecture', level=3)
    doc.add_paragraph('The component hierarchy follows atomic design principles:')
    components = [
        'Atoms: Base UI primitives (buttons, inputs, badges) from Motif design system',
        'Molecules: Composed components (form groups, card widgets, metric displays)',
        'Organisms: Complex components (data tables, chart panels, navigation menus)',
        'Templates: Page layouts (sidebar navigation, content areas, header)',
        'Pages: Complete views combining organisms and templates',
    ]
    for item in components:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('3.4.2 State Management Strategy', level=3)
    doc.add_paragraph('The application employs a dual state management approach:')
    
    state_mgmt = [
        ('Server State:', 'TanStack React Query v5 manages all data fetched from the backend. React Query provides caching, background refetching, optimistic updates, and automatic retry logic.'),
        ('Client State:', 'React Context API manages authentication state, user preferences, and UI state that does not originate from the server.'),
    ]
    for title, desc in state_mgmt:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    # 3.5
    doc.add_heading('3.5 Database Design', level=2)
    doc.add_paragraph('The database design employs polyglot persistence, using the most appropriate database technology for each data type:')
    
    doc.add_heading('3.5.1 PostgreSQL (Primary Transactional Database)', level=3)
    doc.add_paragraph('Core tables include: ai_assets, risk_assessments, compliance_records, vulnerabilities, shadow_ai_detections, users, audit_logs, scan_history, and reports.')
    
    doc.add_heading('3.5.2 Redis (Cache and Rate Limiting)', level=3)
    redis_items = ['Session cache for frequently accessed data', 'Rate limiting counters (sliding window algorithm)', 'Real-time metric aggregation', 'WebSocket pub/sub for real-time updates']
    for item in redis_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('3.5.3 Neo4j (Graph Database)', level=3)
    neo4j_items = ['AI asset relationship graphs', 'Attack path visualization', 'Data lineage tracking', 'Impact analysis']
    for item in neo4j_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('3.5.4 MongoDB (Document Store)', level=3)
    mongo_items = ['Detailed audit log entries', 'Scan result documents', 'Report content storage', 'Configuration history']
    for item in mongo_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 3.6
    doc.add_heading('3.6 Security Architecture', level=2)
    doc.add_paragraph('Security is implemented as a cross-cutting concern with defense-in-depth:')
    
    doc.add_heading('3.6.1 Authentication Flow', level=3)
    auth_flow = [
        'User submits credentials to /api/v1/auth/login',
        'Backend validates credentials against bcrypt-hashed password',
        'If MFA enabled, user must provide TOTP code',
        'On success, backend generates JWT access token (1-hour) and refresh token (7-day)',
        'Access token includes user ID, role, and permissions claims',
        'Frontend stores tokens securely and includes in all API requests',
        'On token expiry, frontend uses refresh token for renewal',
        'On logout, tokens are blacklisted to prevent reuse',
    ]
    for i, item in enumerate(auth_flow, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    doc.add_heading('3.6.2 Authorization Model', level=3)
    doc.add_paragraph('Role-Based Access Control (RBAC) with permission-based route protection. Roles define permission sets, each API endpoint declares required permissions, and middleware validates user permissions against endpoint requirements.')
    
    doc.add_heading('3.6.3 Data Protection', level=3)
    data_prot = [
        'All sensitive fields encrypted at rest using AES-256',
        'All communication encrypted in transit using TLS 1.3',
        'PII detected and masked in logs and non-privileged API responses',
        'Database connection strings and API keys stored in environment variables',
        'Secret rotation supported through secrets manager integration',
    ]
    for item in data_prot:
        doc.add_paragraph(item, style='List Bullet')
    
    # 3.7
    doc.add_heading('3.7 Deployment Architecture', level=2)
    
    doc.add_heading('3.7.1 Development Environment', level=3)
    doc.add_paragraph('Docker Compose orchestrates local development with PostgreSQL 15, Redis 7, Neo4j 5, FastAPI backend with hot reload, and Vite dev server with HMR.')
    
    doc.add_heading('3.7.2 Production Environment', level=3)
    doc.add_paragraph('Kubernetes deployment with Horizontal Pod Autoscaler (2-10 replicas), managed database services, Nginx Ingress Controller with TLS termination, Persistent Volume Claims, Network Policies, Resource quotas, and Liveness/readiness probes.')
    
    doc.add_heading('3.7.3 CI/CD Pipeline', level=3)
    doc.add_paragraph('GitHub Actions workflow with stages: Pull Request (lint, type-check, unit tests, security scan), Merge to main (integration tests, build Docker images), Staging deployment (automated with smoke tests), and Production promotion (manual approval gate, canary deployment).')
    
    doc.add_page_break()


def add_chapter4(doc):
    """Add Chapter 4: Implementation."""
    doc.add_heading('CHAPTER 4: IMPLEMENTATION', level=1)
    
    # 4.1
    doc.add_heading('4.1 Development Environment Setup', level=2)
    doc.add_paragraph('The development environment requires Python 3.11+ with virtual environment, Node.js 18+ with npm, Docker Desktop, Git, and VS Code as the primary IDE.')
    doc.add_paragraph('The backend uses 157 production dependencies. Key categories include web framework (FastAPI, Uvicorn, Pydantic), database (SQLAlchemy, asyncpg, Alembic), authentication (python-jose, passlib, bcrypt), ML tools (scikit-learn, evidently), and security (cryptography, presidio).')
    doc.add_paragraph('The frontend uses React 18, TypeScript, TanStack React Query, React Router DOM, Recharts, Axios, and the EY Motif Design System.')
    
    # 4.2
    doc.add_heading('4.2 Backend Implementation', level=2)
    
    doc.add_heading('4.2.1 FastAPI Application Configuration', level=3)
    doc.add_paragraph('The main application configures the FastAPI instance with lifespan management (async context manager for startup/shutdown), middleware stack (CORS, CSRF, rate limiting, security headers), and router registration under /api/v1 prefix.')
    
    doc.add_heading('4.2.2 Rate Limiting Implementation', level=3)
    doc.add_paragraph('Rate limiting uses a Redis-backed sliding window algorithm with per-client request counters, configurable windows, graceful degradation when Redis is unavailable, and rate limit headers in all responses. Configuration supports global limits (100 req/min default), per-endpoint limits, burst handling, and client identification.')
    
    doc.add_heading('4.2.3 Asset Discovery Engine Implementation', level=3)
    doc.add_paragraph('The discovery engine implements the Strategy pattern with a central Orchestrator coordinating multiple Scanner implementations:')
    
    scanners = [
        ('Cloud ML Scanner:', 'Uses boto3 (AWS SageMaker), azure-ai-ml (Azure ML), and google-cloud-aiplatform (GCP Vertex AI) to enumerate models, endpoints, and training jobs.'),
        ('Kubernetes Scanner:', 'Scans K8s clusters for ML workloads using pod labels, identifies ML frameworks in container images, detects model volumes.'),
        ('Shadow AI Scanner:', 'Analyzes network logs for connections to known AI service domains, correlates traffic to user identities.'),
    ]
    for title, desc in scanners:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_heading('4.2.4 Risk Assessment Implementation', level=3)
    doc.add_paragraph('The risk assessment engine implements the FAIR methodology with steps: Threat Event Frequency (TEF), Vulnerability Factor (VF), Loss Event Frequency (LEF = TEF \u00d7 VF), Primary Loss Magnitude, Secondary Loss Magnitude, and composite Risk Score (0-100).')
    
    doc.add_heading('4.2.5 Compliance Engine Implementation', level=3)
    doc.add_paragraph('The compliance engine maintains rule sets for EU AI Act (risk classification, article mapping), NIST AI RMF (function mapping, maturity scoring), and ISO 42001 (clause mapping, control assessment). Assessment workflow includes asset selection, metadata retrieval, rule evaluation, gap identification, recommendation generation, and notification.')
    
    # 4.3
    doc.add_heading('4.3 Frontend Implementation', level=2)
    
    doc.add_heading('4.3.1 Application Bootstrap', level=3)
    doc.add_paragraph('The application initializes with React DOM root, React Router BrowserRouter, React Query QueryClientProvider, and global CSS imports.')
    
    doc.add_heading('4.3.2 Layout Component', level=3)
    doc.add_paragraph('Provides collapsible sidebar navigation with 7 platform categories: Model Insight, Exposure Management, LLM Scanner, Responsible AI, Data Risk, Red Teaming, and Agentic Remediation.')
    
    doc.add_heading('4.3.3 API Service Layer', level=3)
    doc.add_paragraph('Centralized Axios instance with request interceptor (JWT token attachment, correlation ID) and response interceptor (401 handling, error transformation).')
    
    doc.add_heading('4.3.4 Dashboard Implementation', level=3)
    doc.add_paragraph('The Dashboard page aggregates: Risk Score Overview (donut chart), Asset Inventory (bar chart), Compliance Status (progress indicators), Vulnerability Trends (line chart), Shadow AI Activity (real-time feed), Recent Scans (table), and Action Items (prioritized tasks).')
    
    # 4.4
    doc.add_heading('4.4 Core Module Implementation', level=2)
    
    doc.add_heading('4.4.1 Shadow AI Detection Module', level=3)
    doc.add_paragraph('Combines network analysis (Zeek connection logs, AI service domain detection) and endpoint analysis (osquery for installed AI tools, browser extensions). Outputs include service identification, user attribution, usage frequency, risk classification, and governance recommendations.')
    
    doc.add_heading('4.4.2 Vulnerability Management Module', level=3)
    doc.add_paragraph('Integrates MITRE ATLAS (full taxonomy, asset mapping, detection rules), OWASP LLM (assessment criteria, automated scanning), and ML-specific scanning (framework CVE checking, model file integrity, dependency chain analysis).')
    
    doc.add_heading('4.4.3 Responsible AI Module', level=3)
    doc.add_paragraph('Includes Explainability Service (SHAP integration for feature importance and prediction explanations), Fairness Service (AIF360 for bias metrics and disparate impact analysis), and Data Quality Service (Great Expectations for validation and anomaly detection).')
    
    doc.add_heading('4.4.4 Executive Dashboard Module', level=3)
    doc.add_paragraph('Role-based views for CISO (security posture), CTO (technical debt), Chief AI Officer (AI portfolio), Compliance Officer (framework compliance), and Board (high-level risk summary).')
    
    # 4.5
    doc.add_heading('4.5 LLM Safety and Red Teaming Implementation', level=2)
    
    doc.add_heading('4.5.1 Prompt Injection Testing', level=3)
    doc.add_paragraph('Attack vectors include direct injection, indirect injection, context window manipulation, and encoding attacks. Methodology: define baseline, apply attack vectors, evaluate deviations, score vulnerability (0-10), generate reports.')
    
    doc.add_heading('4.5.2 Hallucination Testing', level=3)
    doc.add_paragraph('Test types: Factual Accuracy, Consistency, Citation Verification, Boundary Testing, and Confabulation Detection. Scoring includes accuracy rate, consistency score, citation validity rate, and confabulation rate.')
    
    doc.add_heading('4.5.3 Toxicity Testing', level=3)
    doc.add_paragraph('Categories: Hate Speech, Violence, Sexual Content, Self-Harm, Misinformation, and Bias. Uses adversarial prompts, edge case testing, demographic parity testing, and toxicity classifiers.')
    
    doc.add_heading('4.5.4 Data Leakage Testing', level=3)
    doc.add_paragraph('Test vectors: Training Data Extraction, PII Extraction, Confidential Information, System Prompt Extraction, and Fine-tuning Data Leakage.')
    
    doc.add_heading('4.5.5 Red Team Session Management', level=3)
    doc.add_paragraph('Structured adversarial testing with session configuration, attack tracking, finding management, reporting, and remediation tracking.')
    
    # 4.6
    doc.add_heading('4.6 AI-BOM Generation Implementation', level=2)
    
    doc.add_heading('4.6.1 BOM Generator (630 LOC)', level=3)
    doc.add_paragraph('Produces CycloneDX 1.5 compliant bills of materials with component extraction covering model identity, architecture, training data, framework dependencies, hardware requirements, performance metrics, and security metadata.')
    
    doc.add_heading('4.6.2 Vulnerability Scanner (320 LOC)', level=3)
    doc.add_paragraph('Cross-references BOM components against CVE Database, ML Attack Database, Configuration Checks, and Supply Chain Risks. Risk scoring uses base severity with multipliers for exposure, sensitivity, and user count, with mitigations reducing the score.')
    
    doc.add_heading('4.6.3 Compliance Mapper', level=3)
    doc.add_paragraph('Maps BOM components to NIST AI RMF subcategories, EU AI Act articles, ISO 42001 clauses, and SPDX 3.0 license compatibility.')
    
    # 4.7
    doc.add_heading('4.7 Authentication and Authorization Implementation', level=2)
    
    doc.add_heading('4.7.1 JWT Token Generation', level=3)
    doc.add_paragraph('Token creation using python-jose with HMAC-SHA256 signing. Access token: 1-hour expiry. Refresh token: 7-day expiry. Payload includes user_id, role, permissions, issued_at, and expires_at.')
    
    doc.add_heading('4.7.2 Auth Middleware (449 LOC)', level=3)
    doc.add_paragraph('Request flow: Extract Bearer token, decode and validate JWT, check expiration, extract claims, verify permissions, create authenticated context, log access event, return 401/403 for failures.')
    
    doc.add_heading('4.7.3 Permission System', level=3)
    doc.add_paragraph('Granular permissions as strings (e.g., "asset:read"), roles map to permission sets, endpoints declare required permissions via dependency injection, supports AND/OR logic, admin has all permissions implicitly.')
    
    doc.add_heading('4.7.4 Audit Logging', level=3)
    doc.add_paragraph('All authenticated actions logged with timestamp (UTC), user ID/role, action, resource affected, request details, response status, and duration. Logs are immutable (append-only).')
    
    doc.add_page_break()


def add_chapter5(doc):
    """Add Chapter 5: Testing and Results."""
    doc.add_heading('CHAPTER 5: TESTING AND RESULTS', level=1)
    
    # 5.1
    doc.add_heading('5.1 Testing Strategy', level=2)
    doc.add_paragraph('The testing strategy follows the test pyramid principle: Unit Tests (base), Integration Tests (middle), End-to-End Tests (top), Security Tests (parallel), and Performance Tests (parallel).')
    doc.add_paragraph('Test coverage targets: Backend 70%+ line coverage (achieved), Frontend 60%+ component coverage (achieved), Critical paths 100% coverage for authentication, authorization, and data access.')
    
    # 5.2
    doc.add_heading('5.2 Unit Testing', level=2)
    doc.add_heading('5.2.1 Backend Unit Tests', level=3)
    doc.add_paragraph('Framework: pytest with pytest-asyncio. Categories: Service Logic, Model Validation, Utility Functions, Configuration. Coverage areas include risk score calculation, permission checking, token lifecycle, scanner normalization, and BOM generation.')
    
    doc.add_heading('5.2.2 Frontend Unit Tests', level=3)
    doc.add_paragraph('Framework: Vitest with React Testing Library. Categories: Component Rendering, User Interaction, Hook Behavior, Utility Functions.')
    
    # 5.3
    doc.add_heading('5.3 Integration Testing', level=2)
    doc.add_heading('5.3.1 API Integration Tests', level=3)
    doc.add_paragraph('Tests verify complete request-response cycles: authentication endpoints, CRUD operations, authorization enforcement, error handling, pagination, and filtering.')
    
    doc.add_heading('5.3.2 Cross-Module Integration', level=3)
    doc.add_paragraph('Tests verify: Discovery \u2192 Asset Inventory, Risk Assessment \u2192 Compliance, Vulnerability \u2192 Notification, Auth \u2192 Audit Log.')
    
    # 5.4
    doc.add_heading('5.4 Security Testing', level=2)
    doc.add_heading('5.4.1 Authentication Security Tests', level=3)
    doc.add_paragraph('Tests cover JWT manipulation, token expiration, brute force prevention, invalid signatures, token reuse after logout, and privilege escalation detection.')
    
    doc.add_heading('5.4.2 Input Validation Tests', level=3)
    doc.add_paragraph('Tests cover SQL injection prevention, XSS payload escaping, path traversal blocking, oversized payload rejection, and malformed JSON handling.')
    
    doc.add_heading('5.4.3 OWASP Top 10 Coverage', level=3)
    doc.add_paragraph('Full coverage: A01 Broken Access Control, A02 Cryptographic Failures, A03 Injection, A04 Insecure Design, A05 Security Misconfiguration, A06 Vulnerable Components, A07 Authentication Failures, A08 Data Integrity Failures, A09 Logging Failures, A10 SSRF.')
    
    # 5.5
    doc.add_heading('5.5 Performance Testing', level=2)
    doc.add_heading('5.5.1 Load Testing (k6)', level=3)
    doc.add_paragraph('Scenarios: Baseline (10 users, 5 min), Normal load (100 users, 15 min), Peak load (500 users, 10 min), Stress test (ramp to 1000 users).')
    
    doc.add_heading('5.5.2 Performance Metrics', level=3)
    doc.add_paragraph('Metrics measured: Response time (p50, p95, p99), Throughput (req/s), Error rate, Database query time, Memory usage, CPU utilization.')
    
    # 5.6
    doc.add_heading('5.6 Results and Analysis', level=2)
    
    doc.add_heading('5.6.1 Test Coverage Results', level=3)
    
    # Backend coverage table
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = 'Module'
    cells[1].text = 'Coverage'
    for cell in cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    
    data = [
        ('Overall', '72%'),
        ('Authentication', '95%'),
        ('Risk Assessment', '78%'),
        ('Compliance Engine', '71%'),
        ('Discovery Engine', '68%'),
        ('API Routes', '85%'),
    ]
    for i, (module, coverage) in enumerate(data, 1):
        table.rows[i].cells[0].text = module
        table.rows[i].cells[1].text = coverage
    
    doc.add_paragraph()
    
    doc.add_heading('5.6.2 Performance Results', level=3)
    
    # Performance table
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = 'Metric'
    cells[1].text = 'Normal Load (100 users)'
    cells[2].text = 'Peak Load (500 users)'
    for cell in cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    
    perf_data = [
        ('Response time (p50)', '45ms', '89ms'),
        ('Response time (p95)', '142ms', '198ms'),
        ('Response time (p99)', '287ms', '412ms'),
        ('Throughput', '850 req/s', '2,100 req/s'),
        ('Error rate', '0.02%', '0.15%'),
        ('DB query time (avg)', '12ms', '-'),
        ('Token validation', '1.4ms', '-'),
    ]
    for i, (metric, normal, peak) in enumerate(perf_data, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = normal
        table.rows[i].cells[2].text = peak
    
    doc.add_paragraph()
    
    doc.add_heading('5.6.3 Security Test Results', level=3)
    sec_results = [
        'Zero critical vulnerabilities found in security testing',
        'Zero SQL injection vectors identified',
        'Zero authentication bypass vulnerabilities',
        'All OWASP Top 10 controls verified functional',
        'Rate limiting effectively prevents brute force (verified at 100 req/min)',
        'All sensitive data confirmed encrypted at rest',
        'No PII leakage in API responses or logs',
    ]
    for item in sec_results:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.6.4 Key Findings', level=3)
    findings = [
        'The modular architecture enables effective isolation of concerns, with each module independently testable and deployable.',
        'Async FastAPI with connection pooling delivers excellent performance under load, meeting the <200ms p95 target even at 500 concurrent users.',
        'Redis-backed rate limiting provides consistent protection across multiple application instances without coordination overhead.',
        'The JWT + RBAC authentication system adds minimal overhead (< 2ms) while providing comprehensive access control.',
        'The polyglot persistence strategy appropriately leverages each database technology\u2019s strengths.',
    ]
    for i, finding in enumerate(findings, 1):
        doc.add_paragraph(f'{i}. {finding}')
    
    doc.add_page_break()


def add_chapter6(doc):
    """Add Chapter 6: Conclusion and Future Work."""
    doc.add_heading('CHAPTER 6: CONCLUSION AND FUTURE WORK', level=1)
    
    # 6.1
    doc.add_heading('6.1 Summary of Contributions', level=2)
    doc.add_paragraph('This project has delivered a comprehensive AI Security Posture Management platform that addresses the critical gap in enterprise AI security tooling. The key contributions include:')
    
    contributions = [
        'Unified AISPM Architecture: Design and implementation of a modular platform integrating asset discovery, risk assessment, compliance monitoring, vulnerability management, LLM safety testing, and governance automation.',
        'Multi-Source Asset Discovery: Implementation of a 10-scanner discovery engine capable of identifying AI assets across cloud platforms, Kubernetes clusters, code repositories, file systems, MLOps platforms, and network traffic.',
        'AI-BOM Generation: Development of an industry-first AI Bill of Materials generator compliant with CycloneDX 1.5, enabling supply chain transparency for AI systems.',
        'LLM Safety Framework: Implementation of a structured LLM security testing framework covering prompt injection, hallucination, toxicity, and data leakage.',
        'Open-Source Integration: Orchestration of 14 industry-leading open-source tools into a cohesive platform.',
        'Enterprise-Grade Security: Comprehensive security controls including JWT authentication, RBAC, distributed rate limiting, CSRF protection, and AES-256 encryption.',
        'Production-Ready Deployment: Complete Docker, Kubernetes, and CI/CD configurations.',
    ]
    for i, item in enumerate(contributions, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    # 6.2
    doc.add_heading('6.2 Key Achievements', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Quantitative Achievements:')
    run.bold = True
    
    quant = [
        '26,000+ lines of production code delivered',
        '70+ backend service modules implemented',
        '28+ frontend pages with interactive dashboards',
        '60+ API endpoints with comprehensive documentation',
        '12 core modules covering the complete AISPM lifecycle',
        '14 open-source tools integrated',
        '7 user roles with 20+ granular permissions',
        '5+ regulatory frameworks supported',
        '25+ test files with 70%+ backend coverage',
        'Sub-200ms API response times at p95',
    ]
    for item in quant:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('Qualitative Achievements:')
    run.bold = True
    
    qual = [
        'Addresses all 6 identified research gaps from literature review',
        'Provides capabilities not available in any single competing solution',
        'Demonstrates feasibility of unified AISPM platform architecture',
        'Comprehensive documentation (150+ pages)',
        'Follows enterprise software engineering best practices',
    ]
    for item in qual:
        doc.add_paragraph(item, style='List Bullet')
    
    # 6.3
    doc.add_heading('6.3 Limitations', level=2)
    limitations = [
        'Deployment Scope: Platform developed and tested in isolated environments rather than deployed against live enterprise AI infrastructure.',
        'Scanner Coverage: While 10 scanner types are implemented, coverage of all possible AI deployment patterns is not complete.',
        'Compliance Depth: Compliance rules represent an interpretation of regulatory text not validated by legal counsel.',
        'LLM Testing Scope: Safety testing covers major categories but not all possible attack variations.',
        'Performance at Scale: Testing conducted with synthetic data; characteristics with 10,000+ assets not validated.',
        'Integration Depth: Open-source tool integrations focus on primary capabilities.',
    ]
    for i, item in enumerate(limitations, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    # 6.4
    doc.add_heading('6.4 Future Work', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Near-Term (1-3 months):')
    run.bold = True
    near = ['Production deployment to enterprise environment', 'Additional scanner types', 'Enhanced LLM safety test suites', 'Mobile application completion', 'ML-powered anomaly detection']
    for item in near:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('Medium-Term (3-6 months):')
    run.bold = True
    med = ['Full SaaS multi-tenant deployment', 'Marketplace for community integrations', 'Automated compliance evidence generation', 'Advanced attack simulation', 'Real-time model behavior monitoring']
    for item in med:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('Long-Term (6-12 months):')
    run.bold = True
    long_term = ['Quantum-resistant encryption', 'AI-powered code review for ML pipelines', 'Predictive remediation', 'Industry-specific compliance modules', 'Agentic AI governance with autonomous remediation']
    for item in long_term:
        doc.add_paragraph(item, style='List Bullet')
    
    # 6.5
    doc.add_heading('6.5 Lessons Learned', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Technical Lessons:')
    run.bold = True
    tech = [
        'Modular monolith architecture provides the right balance of modularity and operational simplicity.',
        'Polyglot persistence pays dividends when data access patterns are clearly different.',
        'FastAPI\u2019s async capabilities enable excellent performance without thread-based complexity.',
        'TypeScript\u2019s type system catches entire categories of frontend bugs at compile time.',
        'React Query eliminates significant boilerplate for data fetching and caching.',
    ]
    for i, item in enumerate(tech, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    p = doc.add_paragraph()
    run = p.add_run('Professional Lessons:')
    run.bold = True
    prof = [
        'Enterprise software requires attention to non-functional requirements from the start.',
        'Documentation is a first-class deliverable in enterprise contexts.',
        'Security must be integrated into every layer of the application.',
        'The AI security domain is rapidly evolving, requiring adaptable architecture.',
        'Open-source tools require engineering effort to integrate and operationalize.',
    ]
    for i, item in enumerate(prof, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    doc.add_page_break()


def add_references(doc):
    """Add References section."""
    doc.add_heading('REFERENCES', level=1)
    
    refs = [
        'Biggio, B., & Roli, F. (2018). Wild patterns: Ten years after the rise of adversarial machine learning. Pattern Recognition, 84, 317-331.',
        'Carlini, N., & Wagner, D. (2017). Towards evaluating the robustness of neural networks. In 2017 IEEE Symposium on Security and Privacy (SP) (pp. 39-57). IEEE.',
        'Carlini, N., Ippolito, D., Jagielski, M., Lee, K., Tramer, F., & Zhang, C. (2023). Quantifying memorization across neural language models. In International Conference on Learning Representations (ICLR 2023).',
        'Chen, P. Y., Zhang, H., Sharma, Y., Yi, J., & Hsieh, C. J. (2017). ZOO: Zeroth order optimization based black-box attacks to deep neural networks without training substitute models. In Proceedings of the 10th ACM Workshop on Artificial Intelligence and Security (pp. 15-26).',
        'European Parliament. (2024). Regulation (EU) 2024/1689 laying down harmonised rules on artificial intelligence (Artificial Intelligence Act). Official Journal of the European Union.',
        'Fredrikson, M., Jha, S., & Ristenpart, T. (2015). Model inversion attacks that exploit confidence information and basic countermeasures. In Proceedings of the 22nd ACM SIGSAC Conference (pp. 1322-1333).',
        'Goldblum, M., et al. (2022). Dataset security for machine learning: Data poisoning, backdoor attacks, and defenses. IEEE Transactions on Pattern Analysis and Machine Intelligence, 45(2), 1563-1580.',
        'Goodfellow, I. J., Shlens, J., & Szegedy, C. (2015). Explaining and harnessing adversarial examples. In International Conference on Learning Representations (ICLR 2015).',
        'Gu, T., Liu, K., Dolan-Gavitt, B., & Garg, S. (2019). BadNets: Evaluating backdooring attacks on deep neural networks. IEEE Access, 7, 47230-47244.',
        'International Organization for Standardization. (2023). ISO/IEC 42001:2023 Information technology - Artificial intelligence - Management system. ISO.',
        'MITRE Corporation. (2023). ATLAS: Adversarial Threat Landscape for Artificial-Intelligence Systems. https://atlas.mitre.org/',
        'National Institute of Standards and Technology. (2023). Artificial Intelligence Risk Management Framework (AI RMF 1.0). NIST AI 100-1.',
        'OWASP Foundation. (2023). OWASP Top 10 for Large Language Model Applications.',
        'Perez, F., & Ribeiro, I. (2022). Ignore this title and HackAPrompt: Exposing systemic weaknesses of LLMs. In Proceedings of EMNLP 2022.',
        'Sculley, D., et al. (2015). Hidden technical debt in machine learning systems. In Advances in Neural Information Processing Systems 28 (pp. 2503-2511).',
        'Shokri, R., Stronati, M., Song, C., & Shmatikov, V. (2017). Membership inference attacks against machine learning models. In 2017 IEEE S&P (pp. 3-18).',
        'Tramer, F., Zhang, F., Juels, A., Reiter, M. K., & Ristenpart, T. (2016). Stealing machine learning models via prediction APIs. In 25th USENIX Security Symposium (pp. 601-618).',
        'FastAPI. (2024). FastAPI framework documentation. https://fastapi.tiangolo.com/',
        'React. (2024). React: A JavaScript library for building user interfaces. https://react.dev/',
        'SQLAlchemy. (2024). SQLAlchemy: The Python SQL toolkit and ORM. https://www.sqlalchemy.org/',
        'OWASP CycloneDX. (2024). CycloneDX specification. https://cyclonedx.org/',
        'IBM Research. (2024). Adversarial Robustness Toolbox. https://github.com/Trusted-AI/adversarial-robustness-toolbox',
        'IBM Research. (2024). AI Fairness 360. https://github.com/Trusted-AI/AIF360',
        'Microsoft. (2024). Presidio: Data protection and de-identification SDK. https://github.com/microsoft/presidio',
        'Flower Labs. (2024). Flower: A friendly federated learning framework. https://flower.ai/',
        'OpenDP. (2024). OpenDP: The open-source differential privacy library. https://opendp.org/',
        'Trail of Bits. (2024). Fickling: A Python pickling decompiler. https://github.com/trailofbits/fickling',
        'McKinsey & Company. (2025). The state of AI in 2025: Global survey results. McKinsey Digital.',
        'Gartner. (2025). Market guide for AI security posture management. Gartner Research.',
    ]
    
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    
    doc.add_page_break()


def add_appendix_a(doc):
    """Add Appendix A: API Endpoint Reference."""
    doc.add_heading('APPENDIX A: API ENDPOINT REFERENCE', level=1)
    
    sections = {
        'Authentication': [
            ('POST', '/api/v1/auth/login', 'Authenticate user, return JWT tokens'),
            ('POST', '/api/v1/auth/register', 'Create new user account'),
            ('POST', '/api/v1/auth/refresh', 'Refresh expired access token'),
            ('POST', '/api/v1/auth/logout', 'Invalidate current tokens'),
            ('GET', '/api/v1/auth/me', 'Get current user profile'),
            ('PUT', '/api/v1/auth/me', 'Update user profile'),
            ('POST', '/api/v1/auth/mfa/setup', 'Configure MFA'),
            ('POST', '/api/v1/auth/mfa/verify', 'Verify MFA code'),
        ],
        'Asset Management': [
            ('GET', '/api/v1/assets', 'List all assets (paginated)'),
            ('POST', '/api/v1/assets', 'Create new asset record'),
            ('GET', '/api/v1/assets/{id}', 'Get asset details'),
            ('PUT', '/api/v1/assets/{id}', 'Update asset record'),
            ('DELETE', '/api/v1/assets/{id}', 'Delete asset record'),
            ('POST', '/api/v1/assets/discover', 'Trigger discovery scan'),
            ('GET', '/api/v1/assets/statistics', 'Get portfolio statistics'),
        ],
        'Risk Management': [
            ('GET', '/api/v1/risks', 'List risk assessments'),
            ('POST', '/api/v1/risks/assess', 'Trigger risk assessment'),
            ('GET', '/api/v1/risks/{id}', 'Get assessment details'),
            ('GET', '/api/v1/risks/dashboard', 'Get risk dashboard'),
            ('GET', '/api/v1/risks/trends', 'Get risk trends'),
        ],
        'LLM Safety': [
            ('POST', '/api/v1/llm-safety/test', 'Run LLM safety test'),
            ('GET', '/api/v1/llm-safety/results', 'Get test results'),
            ('POST', '/api/v1/llm-safety/red-team', 'Start red team session'),
            ('GET', '/api/v1/llm-safety/sessions', 'List red team sessions'),
        ],
        'AI-BOM': [
            ('POST', '/api/v1/ai-bom/generate', 'Generate AI-BOM'),
            ('GET', '/api/v1/ai-bom/{id}', 'Get AI-BOM by ID'),
            ('GET', '/api/v1/ai-bom/list', 'List all AI-BOMs'),
            ('POST', '/api/v1/ai-bom/{id}/export', 'Export BOM'),
            ('POST', '/api/v1/ai-bom/compare', 'Compare two BOMs'),
        ],
    }
    
    for section_name, endpoints in sections.items():
        p = doc.add_paragraph()
        run = p.add_run(section_name)
        run.bold = True
        
        table = doc.add_table(rows=len(endpoints) + 1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Method'
        hdr[1].text = 'Endpoint'
        hdr[2].text = 'Description'
        for cell in hdr:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.bold = True
        
        for i, (method, path, desc) in enumerate(endpoints, 1):
            table.rows[i].cells[0].text = method
            table.rows[i].cells[1].text = path
            table.rows[i].cells[2].text = desc
        
        doc.add_paragraph()
    
    doc.add_page_break()


def add_appendix_b(doc):
    """Add Appendix B: Database Schema."""
    doc.add_heading('APPENDIX B: DATABASE SCHEMA', level=1)
    
    tables_info = [
        ('ai_assets', [
            ('id', 'UUID PRIMARY KEY'),
            ('name', 'VARCHAR(255) NOT NULL'),
            ('asset_type', 'VARCHAR(50) NOT NULL'),
            ('status', 'VARCHAR(50) DEFAULT \'ACTIVE\''),
            ('risk_tier', 'VARCHAR(20) DEFAULT \'MEDIUM\''),
            ('risk_score', 'FLOAT DEFAULT 0.0'),
            ('framework', 'VARCHAR(100)'),
            ('owner', 'VARCHAR(255)'),
            ('department', 'VARCHAR(255)'),
            ('deployment_environment', 'VARCHAR(50)'),
            ('cloud_provider', 'VARCHAR(50)'),
            ('data_classification', 'VARCHAR(50)'),
            ('vulnerability_count', 'INTEGER DEFAULT 0'),
            ('compliance_frameworks', 'JSONB'),
            ('created_at', 'TIMESTAMP DEFAULT NOW()'),
            ('updated_at', 'TIMESTAMP DEFAULT NOW()'),
        ]),
        ('users', [
            ('id', 'UUID PRIMARY KEY'),
            ('email', 'VARCHAR(255) UNIQUE NOT NULL'),
            ('username', 'VARCHAR(100) UNIQUE NOT NULL'),
            ('password_hash', 'VARCHAR(255) NOT NULL'),
            ('role', 'VARCHAR(50) NOT NULL'),
            ('permissions', 'JSONB'),
            ('mfa_enabled', 'BOOLEAN DEFAULT FALSE'),
            ('is_active', 'BOOLEAN DEFAULT TRUE'),
            ('created_at', 'TIMESTAMP DEFAULT NOW()'),
        ]),
        ('vulnerabilities', [
            ('id', 'UUID PRIMARY KEY'),
            ('asset_id', 'UUID REFERENCES ai_assets(id)'),
            ('title', 'VARCHAR(500) NOT NULL'),
            ('severity', 'VARCHAR(20) NOT NULL'),
            ('category', 'VARCHAR(100)'),
            ('atlas_technique', 'VARCHAR(50)'),
            ('owasp_category', 'VARCHAR(50)'),
            ('status', 'VARCHAR(50) DEFAULT \'OPEN\''),
            ('discovered_at', 'TIMESTAMP DEFAULT NOW()'),
        ]),
    ]
    
    for table_name, columns in tables_info:
        p = doc.add_paragraph()
        run = p.add_run(f'TABLE: {table_name}')
        run.bold = True
        
        table = doc.add_table(rows=len(columns) + 1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Column'
        hdr[1].text = 'Type / Constraints'
        for cell in hdr:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.bold = True
        
        for i, (col, typ) in enumerate(columns, 1):
            table.rows[i].cells[0].text = col
            table.rows[i].cells[1].text = typ
        
        doc.add_paragraph()
    
    doc.add_page_break()


def add_appendix_c(doc):
    """Add Appendix C: Configuration Parameters."""
    doc.add_heading('APPENDIX C: CONFIGURATION PARAMETERS', level=1)
    
    configs = {
        'Core Application': [
            ('APP_NAME', 'AI-SPM Platform'),
            ('APP_VERSION', '1.0.0'),
            ('ENVIRONMENT', 'production|staging|development'),
            ('API_V1_PREFIX', '/api/v1'),
        ],
        'Database': [
            ('DATABASE_URL', 'postgresql+asyncpg://user:pass@host:5432/ai_spm'),
            ('DB_POOL_SIZE', '10'),
            ('DB_MAX_OVERFLOW', '20'),
        ],
        'Authentication': [
            ('SECRET_KEY', '<min-32-chars>'),
            ('ACCESS_TOKEN_EXPIRE_MINUTES', '60'),
            ('REFRESH_TOKEN_EXPIRE_DAYS', '7'),
            ('ALGORITHM', 'HS256'),
        ],
        'Security': [
            ('CORS_ORIGINS', '["https://app.domain.com"]'),
            ('RATE_LIMIT_PER_MINUTE', '100'),
            ('MAX_UPLOAD_SIZE_MB', '100'),
        ],
        'Compliance': [
            ('EU_AI_ACT_ENABLED', 'true'),
            ('NIST_AI_RMF_ENABLED', 'true'),
            ('ISO_42001_ENABLED', 'true'),
            ('OWASP_LLM_ENABLED', 'true'),
            ('MITRE_ATLAS_ENABLED', 'true'),
        ],
    }
    
    for section, params in configs.items():
        p = doc.add_paragraph()
        run = p.add_run(section)
        run.bold = True
        
        table = doc.add_table(rows=len(params) + 1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Parameter'
        hdr[1].text = 'Value / Default'
        for cell in hdr:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.bold = True
        
        for i, (param, val) in enumerate(params, 1):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = val
        
        doc.add_paragraph()


def main():
    """Generate the complete VIT thesis DOCX."""
    print("Creating document...")
    doc = Document()
    
    # Setup
    setup_styles(doc)
    set_margins(doc)
    add_page_numbers(doc)
    
    # Front Matter
    print("Adding cover page...")
    add_cover_page(doc)
    add_declaration(doc)
    add_certificates(doc)
    add_abstract(doc)
    add_acknowledgement(doc)
    add_table_of_contents(doc)
    add_list_of_tables(doc)
    add_list_of_figures(doc)
    add_abbreviations(doc)
    
    # Chapters
    print("Adding Chapter 1...")
    add_chapter1(doc)
    print("Adding Chapter 2...")
    add_chapter2(doc)
    print("Adding Chapter 3...")
    add_chapter3(doc)
    print("Adding Chapter 4...")
    add_chapter4(doc)
    print("Adding Chapter 5...")
    add_chapter5(doc)
    print("Adding Chapter 6...")
    add_chapter6(doc)
    
    # Back Matter
    print("Adding References...")
    add_references(doc)
    print("Adding Appendices...")
    add_appendix_a(doc)
    add_appendix_b(doc)
    add_appendix_c(doc)
    
    # Save
    print(f"Saving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done! Report generated successfully.")


if __name__ == '__main__':
    main()
